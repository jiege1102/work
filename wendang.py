# # 导入库
# from docx import Document
#
# # 新建空白文档
# doc_1 = Document()
#
# # 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9）
# doc_1.add_heading('新建空白文档标题，级别为0',level = 0)
# doc_1.add_heading('新建空白文档标题，级别为1',level = 1)
# doc_1.add_heading('新建空白文档标题，级别为2',level = 2)
#
# # 新增段落
# paragraph_1 = doc_1.add_paragraph('这是第一段文字的开始\n请多多关照！')
# # 加粗
# paragraph_1.add_run('加粗字体').bold = True
# paragraph_1.add_run('普通字体')
# # 斜体
# paragraph_1.add_run('斜体字体').italic =True
#
# # 新段落（当前段落的下方）
# paragraph_2 = doc_1.add_paragraph('新起的第二段文字。')
#
# # 新段落（指定端的上方）
# prior_paragraph = paragraph_1.insert_paragraph_before('在第一段文字前插入的段落')
#
# # 添加分页符(可以进行灵活的排版）
# doc_1.add_page_break()
# # 新段落（指定端的上方）
# paragraph_3 = doc_1.add_paragraph('这是第二页第一段文字！')
#
# # 保存文件（当前目录下）
# doc_1.save('doc_1.docx')



from docx import Document
from docx.shared import RGBColor, Pt,Inches,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# 新建文档（Datawhale）
doc_1 = Document()

# 字体设置（全局）
'''只更改font.name是不够的，还需要调用._element.rPr.rFonts的set()方法。'''

doc_1.styles['Normal'].font.name = u'宋体'
doc_1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9，0时候自动带下划线）
#WD_ALIGN_PARAGRAPH. LEFT：左对齐；
#WD_ALIGN_PARAGRAPH. CENTER：居中对其；
#WD_ALIGN_PARAGRAPH. RIGHT：右对齐；
#WD_ALIGN_PARAGRAPH. JUSTIFY：两端对齐；
heading_1 = doc_1.add_heading('周杰伦',level = 0)
heading_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   #居中对齐，默认居左对齐

# 新增段落
paragraph_1 = doc_1.add_paragraph()
'''
设置段落格式：首行缩进0.75cm，居左，段后距离1.0英寸,1.5倍行距。
'''
paragraph_1.paragraph_format.first_line_indent = Cm(0.75)
paragraph_1.paragraph_format.alignment =  WD_PARAGRAPH_ALIGNMENT.LEFT
paragraph_1.paragraph_format.space_after =  Inches(1.0)
paragraph_1.paragraph_format.line_spacing =  1.5

text = '中国台湾华语流行歌手、' \
       '音乐创作家、作曲家、作词人、' \
       '制作人、杰威尔音乐公司老板之一、导演。' \
       '近年涉足电影行业。周杰伦是2000年后亚洲流行乐坛最具革命性与指标' \
       '性的创作歌手，有“亚洲流行天王”之称。他突破原有亚洲音乐的主题、形' \
       '式，融合多元的音乐素材，创造出多变的歌曲风格，尤以融合中西式曲风的嘻哈' \
       '或节奏蓝调最为著名，可说是开创华语流行音乐“中国风”的先声。周杰伦的' \
       '出现打破了亚洲流行乐坛长年停滞不前的局面，为亚洲流行乐坛翻开了新的一页！'

r_1 = paragraph_1.add_run(text)
r_1.font.size =Pt(10)    #字号
r_1.font.bold =True       #加粗
r_1.font.color.rgb =RGBColor(255,0,0)      #颜色

print(len(paragraph_1.runs))    # 查看段落拥有的run对象数量
print(paragraph_1.runs[0].text)  # 查看对应run对象的文本等属性

# 保存文件（当前目录下）
doc_1.save('周杰伦.docx')